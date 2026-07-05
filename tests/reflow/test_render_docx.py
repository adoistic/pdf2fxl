from docx import Document
from pdf2fxl.reflow.render_docx import render_docx
from pdf2fxl.reflow.docmodel import Doc, Heading, Paragraph, Run, ChapterBreak


def test_docx_has_heading_and_paragraph(tmp_path):
    doc = Doc(title="T", language="en", nodes=[
        ChapterBreak(),
        Heading(level=1, text="Chapter One"),
        Paragraph(runs=[Run(text="Body text here.")]),
    ])
    out = tmp_path / "out.docx"
    render_docx(doc, out, assets_root=tmp_path)
    assert out.exists()
    d = Document(str(out))
    texts = [p.text for p in d.paragraphs]
    assert "Chapter One" in texts
    assert "Body text here." in texts
    styles = [p.style.name for p in d.paragraphs if p.text == "Chapter One"]
    assert styles and styles[0].startswith("Heading")
