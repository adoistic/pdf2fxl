from __future__ import annotations
from pathlib import Path

from docx import Document
from docx.shared import Inches

from .docmodel import Doc, Heading, Paragraph, Figure, Table, Formula, ChapterBreak


def render_docx(doc: Doc, out_path: Path, assets_root: Path) -> Path:
    d = Document()
    d.core_properties.title = doc.title
    for n in doc.nodes:
        if isinstance(n, ChapterBreak):
            if len(d.paragraphs) > 1:
                d.add_page_break()
        elif isinstance(n, Heading):
            d.add_heading(n.text, level=min(9, max(1, n.level)))
        elif isinstance(n, Paragraph):
            p = d.add_paragraph()
            for r in n.runs:
                run = p.add_run(r.text)
                run.bold = r.bold
                run.italic = r.italic
                run.underline = getattr(r, "underline", False)
        elif isinstance(n, (Figure, Table)) and getattr(n, "image_src", None) or (
                isinstance(n, Figure) and n.src):
            src = n.src if isinstance(n, Figure) else n.image_src
            path = Path(assets_root) / Path(src).name
            if path.exists():
                width = Inches(6.0 * getattr(n, "width_frac", 1.0)) if isinstance(n, Figure) else Inches(6.0)
                d.add_picture(str(path), width=width)
            if getattr(n, "caption", None):
                d.add_paragraph(n.caption)
        elif isinstance(n, Formula) and n.text:
            d.add_paragraph(n.text)
    out_path = Path(out_path)
    d.save(str(out_path))
    return out_path
